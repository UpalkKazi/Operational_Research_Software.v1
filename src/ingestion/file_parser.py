"""
File Parser — extract structured data and raw text from uploaded files.

Supported formats: Excel (.xlsx/.xls), CSV, Word (.docx/.doc),
plain text (.txt), PDF, and MPS (.mps/.mps.gz).
"""

import gzip
import io
import warnings
from pathlib import Path
from typing import Any, Dict, List, Union


class FileParser:
    """
    Parse uploaded files into a structured dict that always includes a
    ``raw_text`` key containing a plain-text representation of the full
    content (suitable for passing to the AI classifier).
    """

    _EXT_HANDLERS = {
        '.xlsx': '_parse_excel',
        '.xls': '_parse_excel',
        '.csv': '_parse_csv',
        '.docx': '_parse_docx',
        '.doc': '_parse_docx',
        '.txt': '_parse_text',
        '.pdf': '_parse_pdf',
        '.mps': '_parse_mps',
        '.mps.gz': '_parse_mps',
    }

    def parse(
        self,
        file_path_or_bytes: Union[str, Path, bytes, io.BytesIO],
        filename: str = '',
    ) -> Dict[str, Any]:
        """
        Detect file type and extract content.

        Args:
            file_path_or_bytes: A file path (str/Path) **or** raw bytes /
                BytesIO (e.g. from ``st.file_uploader``).
            filename: Original filename — used for extension detection when
                *file_path_or_bytes* is bytes.

        Returns:
            A dict that **always** contains at least ``type``, ``raw_text``,
            and (on failure) ``error``.
        """
        try:
            ext = self._detect_extension(file_path_or_bytes, filename)
            handler_name = self._EXT_HANDLERS.get(ext)
            if not handler_name:
                return {
                    'type': 'error',
                    'raw_text': '',
                    'error': (
                        f"Unsupported file type '{ext}'. "
                        f"Supported: {', '.join(sorted(self._EXT_HANDLERS))}"
                    ),
                }

            data = self._to_bytes(file_path_or_bytes)
            if ext.endswith('.gz'):
                data = gzip.decompress(data)
            handler = getattr(self, handler_name)
            result = handler(data, filename)

            if 'raw_text' not in result:
                result['raw_text'] = self._build_raw_text(result)

            if isinstance(file_path_or_bytes, (str, Path)):
                result['file_path'] = str(file_path_or_bytes)

            return result

        except Exception as e:
            return {'type': 'error', 'raw_text': '', 'error': str(e)}

    # ------------------------------------------------------------------
    #  Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _detect_extension(
        source: Union[str, Path, bytes, io.BytesIO],
        filename: str,
    ) -> str:
        name = ''
        if filename:
            name = filename
        elif isinstance(source, (str, Path)):
            name = str(source)

        if not name:
            return ''

        low = name.lower()
        if low.endswith('.mps.gz'):
            return '.mps.gz'
        return Path(low).suffix

    @staticmethod
    def _to_bytes(source: Union[str, Path, bytes, io.BytesIO]) -> bytes:
        if isinstance(source, bytes):
            return source
        if isinstance(source, io.BytesIO):
            source.seek(0)
            return source.read()
        with open(source, 'rb') as f:
            return f.read()

    # ------------------------------------------------------------------
    #  Raw-text builder (fallback when handler doesn't set raw_text)
    # ------------------------------------------------------------------

    @staticmethod
    def _build_raw_text(result: Dict[str, Any]) -> str:
        rtype = result.get('type', '')

        if rtype == 'excel':
            parts: List[str] = []
            for sheet in result.get('sheets', []):
                parts.append(f"--- Sheet: {sheet['name']} ---")
                headers = sheet.get('headers', [])
                if headers:
                    parts.append('\t'.join(str(h) for h in headers))
                for row in sheet.get('rows', []):
                    parts.append(
                        '\t'.join(str(row.get(h, '')) for h in headers)
                    )
            return '\n'.join(parts)

        if rtype == 'csv':
            sheets = result.get('sheets', [])
            if sheets:
                sheet = sheets[0]
                headers = sheet.get('headers', [])
                lines = ['\t'.join(str(h) for h in headers)]
                for row in sheet.get('rows', []):
                    lines.append(
                        '\t'.join(str(row.get(h, '')) for h in headers)
                    )
                return '\n'.join(lines)
            return ''

        if rtype == 'docx':
            parts = list(result.get('paragraphs', []))
            for table in result.get('tables', []):
                for row in table:
                    parts.append('\t'.join(str(c) for c in row))
            return '\n'.join(parts)

        if rtype == 'text':
            return result.get('content', '')

        if rtype == 'pdf':
            return '\n\n'.join(result.get('pages', []))

        return ''

    # ------------------------------------------------------------------
    #  Format-specific parsers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_excel(data: bytes, filename: str) -> Dict[str, Any]:
        import openpyxl

        wb = openpyxl.load_workbook(
            io.BytesIO(data), read_only=True, data_only=True,
        )
        sheets: List[Dict[str, Any]] = []
        for name in wb.sheetnames:
            ws = wb[name]
            rows_iter = ws.iter_rows(values_only=True)

            first_row = next(rows_iter, None)
            if first_row is None:
                sheets.append({'name': name, 'headers': [], 'rows': []})
                continue

            headers = [
                str(c) if c is not None else f'col_{i}'
                for i, c in enumerate(first_row)
            ]
            data_rows = [
                {headers[i]: cell for i, cell in enumerate(row)}
                for row in rows_iter
            ]
            sheets.append({
                'name': name,
                'headers': headers,
                'rows': data_rows,
            })

        wb.close()
        return {'type': 'excel', 'sheets': sheets}

    @staticmethod
    def _parse_csv(data: bytes, filename: str) -> Dict[str, Any]:
        import pandas as pd

        text = data.decode('utf-8', errors='replace')
        df = pd.read_csv(io.StringIO(text))
        headers = list(df.columns)
        rows = df.to_dict(orient='records')
        return {
            'type': 'csv',
            'sheets': [{'name': 'Sheet1', 'headers': headers, 'rows': rows}],
        }

    @staticmethod
    def _parse_docx(data: bytes, filename: str) -> Dict[str, Any]:
        import docx

        doc = docx.Document(io.BytesIO(data))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables: List[List[List[str]]] = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables.append(table_data)

        return {'type': 'docx', 'paragraphs': paragraphs, 'tables': tables}

    @staticmethod
    def _parse_text(data: bytes, filename: str) -> Dict[str, Any]:
        content = data.decode('utf-8', errors='replace')
        return {'type': 'text', 'content': content, 'raw_text': content}

    @staticmethod
    def _parse_pdf(data: bytes, filename: str) -> Dict[str, Any]:
        pages: List[str] = []

        # Tier 1: PyMuPDF (fitz)
        try:
            import fitz
            doc = fitz.open(stream=data, filetype='pdf')
            for page in doc:
                pages.append(page.get_text())
            doc.close()
            return {'type': 'pdf', 'pages': pages}
        except ImportError:
            pass

        # Tier 2: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ''
                    pages.append(text)
            return {'type': 'pdf', 'pages': pages}
        except ImportError:
            pass

        raise ImportError(
            "No PDF library available. Install one of: "
            "PyMuPDF (pip install pymupdf) or pdfplumber (pip install pdfplumber)"
        )

    # ------------------------------------------------------------------
    #  MPS (Mathematical Programming System) format
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_mps(data: bytes, filename: str) -> Dict[str, Any]:
        text = data.decode('utf-8', errors='replace')
        lines = text.splitlines()

        problem_name = ''
        objective_name = ''
        rows: List[Dict[str, str]] = []
        columns: Dict[str, Dict[str, float]] = {}
        rhs: Dict[str, float] = {}
        bounds_raw: List[tuple] = []

        _ROW_TYPE_MAP = {
            'N': 'objective',
            'L': '<=',
            'G': '>=',
            'E': '=',
        }

        section = ''
        marker_int = False  # tracks MARKER INT / INTEND blocks

        for lineno, raw_line in enumerate(lines, 1):
            line = raw_line.rstrip()
            if not line:
                continue

            # Section headers start at column 0 (no leading whitespace)
            if line and not line[0].isspace():
                token = line.split()[0].upper()
                if token == 'NAME':
                    section = 'NAME'
                    parts = line.split(None, 1)
                    problem_name = parts[1].strip() if len(parts) > 1 else ''
                    continue
                if token in ('ROWS', 'COLUMNS', 'RHS', 'RANGES',
                             'BOUNDS', 'ENDATA'):
                    section = token
                    continue
                # Unknown section header — skip
                section = token
                continue

            # Data lines (have leading whitespace)
            fields = line.split()
            if not fields:
                continue

            try:
                if section == 'ROWS':
                    rtype = fields[0].upper()
                    rname = fields[1] if len(fields) > 1 else f'row_{lineno}'
                    mapped = _ROW_TYPE_MAP.get(rtype, rtype)
                    rows.append({'name': rname, 'type': mapped})
                    if rtype == 'N':
                        objective_name = rname

                elif section == 'COLUMNS':
                    # Handle MARKER lines for integer variable blocks
                    if len(fields) >= 3 and fields[1].upper() == "'MARKER'":
                        tag = fields[2].strip("'").upper()
                        marker_int = tag == 'INTORG'
                        continue

                    var_name = fields[0]
                    if var_name not in columns:
                        columns[var_name] = {}

                    # One or two (constraint, coef) pairs per line
                    i = 1
                    while i + 1 < len(fields):
                        con_name = fields[i]
                        coef = float(fields[i + 1])
                        columns[var_name][con_name] = coef
                        i += 2

                    if marker_int:
                        bounds_raw.append(('MI', '', var_name, None))

                elif section == 'RHS':
                    # rhs_name con value [con2 value2]
                    i = 1
                    while i + 1 < len(fields):
                        con_name = fields[i]
                        val = float(fields[i + 1])
                        rhs[con_name] = val
                        i += 2

                elif section == 'BOUNDS':
                    btype = fields[0].upper()
                    bnd_name = fields[1] if len(fields) > 1 else ''
                    var_name = fields[2] if len(fields) > 2 else ''
                    val = float(fields[3]) if len(fields) > 3 else None
                    bounds_raw.append((btype, bnd_name, var_name, val))

                elif section == 'RANGES':
                    pass  # RANGES section — not commonly needed

            except Exception as exc:
                warnings.warn(
                    f"MPS parse warning (line {lineno}): {exc} — skipping",
                    stacklevel=2,
                )
                continue

        # --- Build variable_bounds dict -----------------------------------
        all_vars = list(columns.keys())
        variable_bounds: Dict[str, Dict[str, Any]] = {}
        for v in all_vars:
            variable_bounds[v] = {'lb': 0.0, 'ub': None, 'type': 'continuous'}

        for btype, _bnd_name, var_name, val in bounds_raw:
            if var_name not in variable_bounds:
                variable_bounds[var_name] = {
                    'lb': 0.0, 'ub': None, 'type': 'continuous',
                }
            vb = variable_bounds[var_name]

            if btype == 'UP':
                vb['ub'] = val
            elif btype == 'LO':
                vb['lb'] = val
            elif btype == 'FX':
                vb['lb'] = val
                vb['ub'] = val
            elif btype == 'FR':
                vb['lb'] = None
                vb['ub'] = None
            elif btype == 'BV':
                vb['lb'] = 0.0
                vb['ub'] = 1.0
                vb['type'] = 'binary'
            elif btype == 'LI':
                vb['lb'] = val
                vb['type'] = 'integer'
            elif btype == 'UI':
                vb['ub'] = val
                vb['type'] = 'integer'
            elif btype == 'MI':
                vb['type'] = 'integer'

        # --- Separate objective from constraints --------------------------
        obj_coefficients: Dict[str, float] = {}
        constraint_list: List[Dict[str, Any]] = []

        constraint_rows = [r for r in rows if r['type'] != 'objective']
        for row in constraint_rows:
            cname = row['name']
            coeffs: Dict[str, float] = {}
            for v in all_vars:
                if cname in columns.get(v, {}):
                    coeffs[v] = columns[v][cname]
            constraint_list.append({
                'name': cname,
                'type': row['type'],
                'coefficients': coeffs,
                'rhs': rhs.get(cname, 0.0),
            })

        if objective_name:
            for v in all_vars:
                if objective_name in columns.get(v, {}):
                    obj_coefficients[v] = columns[v][objective_name]

        n_vars = len(all_vars)
        n_cons = len(constraint_list)

        # --- Variable type counts ----------------------------------------
        n_binary = sum(1 for vb in variable_bounds.values()
                       if vb['type'] == 'binary')
        n_integer = sum(1 for vb in variable_bounds.values()
                        if vb['type'] == 'integer')
        n_continuous = n_vars - n_binary - n_integer

        # --- Human-readable raw_text for the AI --------------------------
        parts: List[str] = [f"MPS Problem: {problem_name or '(unnamed)'}"]
        type_desc = []
        if n_continuous:
            type_desc.append(f"{n_continuous} continuous")
        if n_integer:
            type_desc.append(f"{n_integer} integer")
        if n_binary:
            type_desc.append(f"{n_binary} binary")
        parts.append(
            f"Variables: {n_vars} ({', '.join(type_desc)})"
            if type_desc else f"Variables: {n_vars}"
        )
        parts.append(f"Constraints: {n_cons}")
        parts.append(f"Objective: minimize {objective_name or '?'}")

        if obj_coefficients:
            top_obj = list(obj_coefficients.items())[:10]
            coef_strs = [f"{v}={c}" for v, c in top_obj]
            suffix = ', ...' if len(obj_coefficients) > 10 else ''
            parts.append(f"Objective coefficients: {', '.join(coef_strs)}{suffix}")

        if constraint_list:
            parts.append("Constraints summary:")
            for con in constraint_list[:5]:
                parts.append(f"  {con['name']} ({con['type']}): rhs={con['rhs']}")
            if n_cons > 5:
                parts.append(f"  ... ({n_cons - 5} more)")

        raw_text = '\n'.join(parts)

        return {
            'type': 'mps',
            'name': problem_name,
            'objective_name': objective_name,
            'objective_sense': 'minimize',
            'variables': all_vars,
            'variable_bounds': variable_bounds,
            'constraints': constraint_list,
            'objective_coefficients': obj_coefficients,
            'num_variables': n_vars,
            'num_constraints': n_cons,
            'raw_text': raw_text,
        }
